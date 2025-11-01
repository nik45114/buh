"""
Генератор документов (договоры, акты)
"""
import logging
from datetime import date
from pathlib import Path
from typing import Optional
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from ..database.models import Employee, Contract
from ..config import settings

logger = logging.getLogger(__name__)


class DocumentGenerator:
    """Генерация документов для сотрудников"""

    def __init__(self, output_dir: str = "/opt/accounting-bot/documents"):
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)
        self.templates_dir = Path("/opt/accounting-bot/templates")
        self.templates_dir.mkdir(parents=True, exist_ok=True)

    def generate_labor_contract(
        self,
        employee: Employee,
        contract: Contract
    ) -> str:
        """
        Генерация трудового договора (ТД)

        Returns:
            Путь к сохраненному файлу
        """
        doc = Document()

        # Настройка стилей
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(12)

        # Заголовок
        heading = doc.add_paragraph()
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = heading.add_run('ТРУДОВОЙ ДОГОВОР')
        run.bold = True
        run.font.size = Pt(14)

        # Номер и дата
        doc.add_paragraph(
            f"№ {contract.contract_number or '_____'} от {contract.start_date.strftime('%d.%m.%Y')} г."
        )
        doc.add_paragraph(f"г. {settings.COMPANY_CITY or 'Москва'}")

        # Стороны договора
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.add_run(f'{settings.COMPANY_NAME}, ИНН {settings.COMPANY_INN}, ')
        p.add_run('именуемое в дальнейшем «Работодатель», в лице директора, ')
        p.add_run('с одной стороны, и ')

        p = doc.add_paragraph()
        p.add_run(f'{employee.full_name}, ')
        if employee.full_passport:
            p.add_run(f'паспорт {employee.full_passport}, ')
        if employee.inn:
            p.add_run(f'ИНН {employee.inn}, ')
        p.add_run('именуемый в дальнейшем «Работник», с другой стороны, ')
        p.add_run('заключили настоящий трудовой договор о нижеследующем:')

        # 1. Предмет договора
        doc.add_paragraph()
        heading = doc.add_paragraph()
        heading.add_run('1. ПРЕДМЕТ ДОГОВОРА').bold = True

        doc.add_paragraph(
            f'1.1. Работник принимается на работу в {settings.COMPANY_NAME} '
            f'на должность {contract.position}.'
        )

        doc.add_paragraph(
            f'1.2. Дата начала работы: {contract.start_date.strftime("%d.%m.%Y")} г.'
        )

        if contract.end_date:
            doc.add_paragraph(
                f'1.3. Договор заключен на определенный срок до {contract.end_date.strftime("%d.%m.%Y")} г.'
            )

        # 2. Оплата труда
        doc.add_paragraph()
        heading = doc.add_paragraph()
        heading.add_run('2. ОПЛАТА ТРУДА').bold = True

        doc.add_paragraph(
            f'2.1. За выполнение трудовых обязанностей Работнику устанавливается '
            f'заработная плата в размере {contract.salary:,.2f} руб. в месяц.'
        )

        doc.add_paragraph(
            '2.2. Заработная плата выплачивается 2 раза в месяц: '
            'аванс - 20 числа текущего месяца, '
            'окончательный расчет - 5 числа следующего месяца.'
        )

        # 3. Обязанности сторон
        doc.add_paragraph()
        heading = doc.add_paragraph()
        heading.add_run('3. ОБЯЗАННОСТИ СТОРОН').bold = True

        doc.add_paragraph('3.1. Работник обязан:')
        doc.add_paragraph('- Добросовестно выполнять свои трудовые обязанности;', style='List Bullet')
        doc.add_paragraph('- Соблюдать трудовую дисциплину;', style='List Bullet')
        doc.add_paragraph('- Бережно относиться к имуществу работодателя.', style='List Bullet')

        doc.add_paragraph('3.2. Работодатель обязан:')
        doc.add_paragraph('- Обеспечить условия труда, предусмотренные ТК РФ;', style='List Bullet')
        doc.add_paragraph('- Своевременно выплачивать заработную плату;', style='List Bullet')
        doc.add_paragraph('- Производить обязательные отчисления в фонды.', style='List Bullet')

        # Подписи
        doc.add_paragraph()
        doc.add_paragraph()
        p = doc.add_paragraph('РАБОТОДАТЕЛЬ:')
        p.add_run('\t\t\t')
        p.add_run('РАБОТНИК:')

        doc.add_paragraph()
        doc.add_paragraph()
        p = doc.add_paragraph(f'{settings.COMPANY_NAME}')
        p.add_run('\t\t\t')
        p.add_run(employee.full_name)

        doc.add_paragraph()
        doc.add_paragraph()
        p = doc.add_paragraph('_______________ / _____________')
        p.add_run('\t\t')
        p.add_run('_______________ / _____________')

        # Сохранение
        filename = f"TD_{employee.id}_{contract.contract_number}_{date.today().isoformat()}.docx"
        filepath = self.output_dir / filename
        doc.save(str(filepath))

        logger.info(f"Labor contract generated: {filepath}")
        return str(filepath)

    def generate_gph_contract(
        self,
        employee: Employee,
        contract: Contract
    ) -> str:
        """
        Генерация договора ГПХ

        Returns:
            Путь к сохраненному файлу
        """
        doc = Document()

        # Настройка стилей
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(12)

        # Заголовок
        heading = doc.add_paragraph()
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = heading.add_run('ДОГОВОР ПОДРЯДА')
        run.bold = True
        run.font.size = Pt(14)

        # Номер и дата
        doc.add_paragraph(
            f"№ {contract.contract_number or '_____'} от {contract.start_date.strftime('%d.%m.%Y')} г."
        )
        doc.add_paragraph(f"г. {settings.COMPANY_CITY or 'Москва'}")

        # Стороны
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.add_run(f'{settings.COMPANY_NAME}, ИНН {settings.COMPANY_INN}, ')
        p.add_run('именуемое в дальнейшем «Заказчик», с одной стороны, и ')

        p = doc.add_paragraph()
        p.add_run(f'{employee.full_name}, ')
        if employee.full_passport:
            p.add_run(f'паспорт {employee.full_passport}, ')
        p.add_run('именуемый в дальнейшем «Исполнитель», с другой стороны, ')
        p.add_run('заключили настоящий договор о нижеследующем:')

        # 1. Предмет договора
        doc.add_paragraph()
        heading = doc.add_paragraph()
        heading.add_run('1. ПРЕДМЕТ ДОГОВОРА').bold = True

        doc.add_paragraph(
            f'1.1. Исполнитель обязуется по заданию Заказчика выполнить работы '
            f'по {contract.work_conditions or "оказанию услуг"}, '
            f'а Заказчик обязуется принять и оплатить выполненные работы.'
        )

        # 2. Стоимость работ
        doc.add_paragraph()
        heading = doc.add_paragraph()
        heading.add_run('2. СТОИМОСТЬ РАБОТ').bold = True

        doc.add_paragraph(
            f'2.1. Стоимость работ по настоящему договору составляет '
            f'{contract.salary:,.2f} руб.'
        )

        # 3. Срок выполнения
        doc.add_paragraph()
        heading = doc.add_paragraph()
        heading.add_run('3. СРОК ВЫПОЛНЕНИЯ РАБОТ').bold = True

        doc.add_paragraph(
            f'3.1. Работы выполняются в период с {contract.start_date.strftime("%d.%m.%Y")} г. '
            f'по {contract.end_date.strftime("%d.%m.%Y") if contract.end_date else "___________"} г.'
        )

        # Подписи
        doc.add_paragraph()
        doc.add_paragraph()
        p = doc.add_paragraph('ЗАКАЗЧИК:')
        p.add_run('\t\t\t')
        p.add_run('ИСПОЛНИТЕЛЬ:')

        doc.add_paragraph()
        doc.add_paragraph()
        p = doc.add_paragraph(f'{settings.COMPANY_NAME}')
        p.add_run('\t\t\t')
        p.add_run(employee.full_name)

        doc.add_paragraph()
        doc.add_paragraph()
        p = doc.add_paragraph('_______________ / _____________')
        p.add_run('\t\t')
        p.add_run('_______________ / _____________')

        # Сохранение
        filename = f"GPH_{employee.id}_{contract.contract_number}_{date.today().isoformat()}.docx"
        filepath = self.output_dir / filename
        doc.save(str(filepath))

        logger.info(f"GPH contract generated: {filepath}")
        return str(filepath)

    def generate_offer(
        self,
        employee: Employee,
        hourly_rate: float,
        position: str = "Администратор"
    ) -> str:
        """
        Генерация договора оферты

        Returns:
            Путь к сохраненному файлу
        """
        doc = Document()

        # Настройка стилей
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(12)

        # Заголовок
        heading = doc.add_paragraph()
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = heading.add_run('ДОГОВОР ОФЕРТЫ')
        run.bold = True
        run.font.size = Pt(14)

        doc.add_paragraph(
            f"на оказание услуг от {date.today().strftime('%d.%m.%Y')} г."
        )
        doc.add_paragraph(f"г. {settings.COMPANY_CITY or 'Москва'}")

        # Текст оферты
        doc.add_paragraph()
        doc.add_paragraph(
            f'{settings.COMPANY_NAME} (далее - «Заказчик») в лице директора, '
            f'предлагает физическим лицам (далее - «Исполнитель») '
            f'заключить договор на следующих условиях:'
        )

        doc.add_paragraph()
        heading = doc.add_paragraph()
        heading.add_run('1. ПРЕДМЕТ ДОГОВОРА').bold = True

        doc.add_paragraph(
            f'1.1. Исполнитель оказывает услуги по должности «{position}» '
            f'в компьютерном клубе Заказчика.'
        )

        doc.add_paragraph()
        heading = doc.add_paragraph()
        heading.add_run('2. УСЛОВИЯ ОПЛАТЫ').bold = True

        doc.add_paragraph(
            f'2.1. Оплата услуг производится по ставке {hourly_rate:,.2f} руб./час.'
        )

        doc.add_paragraph(
            '2.2. Расчет производится по факту отработанного времени в конце месяца.'
        )

        doc.add_paragraph()
        heading = doc.add_paragraph()
        heading.add_run('3. АКЦЕПТ ОФЕРТЫ').bold = True

        doc.add_paragraph(
            '3.1. Фактическое оказание услуг Исполнителем является полным '
            'и безоговорочным акцептом настоящей оферты.'
        )

        # Данные Заказчика
        doc.add_paragraph()
        doc.add_paragraph()
        heading = doc.add_paragraph()
        heading.add_run('ЗАКАЗЧИК:').bold = True

        doc.add_paragraph(settings.COMPANY_NAME)
        doc.add_paragraph(f'ИНН: {settings.COMPANY_INN}')
        doc.add_paragraph(f'Адрес: {getattr(settings, "COMPANY_ADDRESS", "_____")}')

        # Сохранение
        filename = f"OFFER_{employee.id}_{date.today().isoformat()}.docx"
        filepath = self.output_dir / filename
        doc.save(str(filepath))

        logger.info(f"Offer contract generated: {filepath}")
        return str(filepath)
